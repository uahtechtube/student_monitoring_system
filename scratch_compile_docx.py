import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
import re
import os

# Helper to set cell background color
def set_cell_background(cell, fill_hex):
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'
    cell._tc.get_or_add_tcPr().append(parse_xml(shading_xml))

# Helper to set cell borders
def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Format headers
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_background(hdr_cells[i], "4F46E5") # indigo color
        set_cell_margins(hdr_cells[i], top=120, bottom=120, left=150, right=150)
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = docx.shared.RGBColor(255, 255, 255)
            run.font.name = 'Arial'
            run.font.size = Pt(10)

    # Format rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            set_cell_background(row_cells[i], bg_color)
            set_cell_margins(row_cells[i], top=100, bottom=100, left=150, right=150)
            p = row_cells[i].paragraphs[0]
            # If numerical value, right-align or center, else left-align
            if re.match(r'^\s*[\d\.\%\,\-\+x]+\s*$', str(val)):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(10)
    
    # Add an empty paragraph after table
    doc.add_paragraph()

def parse_markdown_to_docx(doc, md_filepath):
    if not os.path.exists(md_filepath):
        print(f"File not found: {md_filepath}")
        return
        
    print(f"Parsing {md_filepath}...")
    with open(md_filepath, "r", encoding="utf-8") as f:
        lines = f.readlines()
        
    in_table = False
    table_headers = []
    table_rows = []
    
    idx = 0
    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()
        
        # Table parsing
        if stripped.startswith("|"):
            if not in_table:
                in_table = True
                # Parse headers
                table_headers = [c.strip() for c in stripped.split("|")[1:-1]]
                # Skip separator line
                idx += 2
                continue
            else:
                # Parse row
                row_data = [c.strip() for c in stripped.split("|")[1:-1]]
                table_rows.append(row_data)
                idx += 1
                continue
        else:
            if in_table:
                add_styled_table(doc, table_headers, table_rows)
                in_table = False
                table_headers = []
                table_rows = []
                
        # Headings
        if stripped.startswith("#### "):
            h = doc.add_heading(level=4)
            r = h.add_run(stripped[5:])
            r.font.name = 'Arial'
            r.font.size = Pt(11)
            r.font.bold = True
            h.paragraph_format.space_before = Pt(6)
            h.paragraph_format.space_after = Pt(4)
        elif stripped.startswith("### "):
            h = doc.add_heading(level=3)
            r = h.add_run(stripped[4:])
            r.font.name = 'Arial'
            r.font.size = Pt(12)
            r.font.bold = True
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            r = h.add_run(stripped[3:])
            r.font.name = 'Arial'
            r.font.size = Pt(14)
            r.font.bold = True
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(8)
        elif stripped.startswith("# "):
            h = doc.add_heading(level=1)
            r = h.add_run(stripped[2:])
            r.font.name = 'Arial'
            r.font.size = Pt(18)
            r.font.bold = True
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.paragraph_format.space_before = Pt(24)
            h.paragraph_format.space_after = Pt(12)
            # Add page break before Chapter headings (if not first chapter)
            if "CHAPTER" in stripped:
                p_break = doc.add_page_break()
        # Bullet Lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            text = stripped[2:]
            # basic bold parser inside list
            parts = text.split("**")
            for part_idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                if part_idx % 2 == 1:
                    run.font.bold = True
            p.paragraph_format.space_after = Pt(4)
        elif stripped:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            
            # Basic bold parsing: **text**
            parts = stripped.split("**")
            for part_idx, part in enumerate(parts):
                run = p.add_run(part)
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                if part_idx % 2 == 1:
                    run.font.bold = True
        idx += 1
        
    if in_table:
        add_styled_table(doc, table_headers, table_rows)

def clean_original_text(txt_filepath):
    print("Reading and cleaning original text...")
    with open(txt_filepath, "r", encoding="utf-8") as f:
        text = f.read()
        
    # Split text into pages
    pages = text.split("=== PAGE")
    cleaned_pages = []
    
    for page in pages:
        if not page.strip():
            continue
        lines = page.split("\n")
        # Strip header/footer lines if they just have numbers
        clean_lines = []
        for line in lines:
            line_str = line.strip()
            # Skip page markers
            if re.match(r'^\d+\s*===', line_str) or not line_str:
                continue
            clean_lines.append(line_str)
        cleaned_pages.append("\n".join(clean_lines))
        
    full_text = "\n\n".join(cleaned_pages)
    return full_text

def compile_all_chapters():
    # 1. Generate Chapters 4 & 5 document
    doc_45 = docx.Document()
    # Set standard margins
    sections = doc_45.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Apply default Normal font
    style = doc_45.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    parse_markdown_to_docx(doc_45, "Chapter_Four.md")
    parse_markdown_to_docx(doc_45, "Chapter_Five.md")
    
    doc_45.save("KIU_Student_Monitoring_Chapters_4_and_5.docx")
    print("Saved KIU_Student_Monitoring_Chapters_4_and_5.docx")
    
    # 2. Generate Full Thesis document (Chapters 1 to 5 + References)
    doc_full = docx.Document()
    for section in doc_full.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    style = doc_full.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Read the original text file
    original_text = clean_original_text("MyProject_text.txt")
    
    # We want to split the original text at "REFERENCES" to insert Chapters 4 and 5 before the references.
    # Let's locate the references header.
    ref_pattern = re.compile(r'\bREFERENCES\b', re.IGNORECASE)
    parts = ref_pattern.split(original_text, maxsplit=1)
    
    chapters_1_3 = parts[0]
    references_text = parts[1] if len(parts) > 1 else ""
    
    # Add Chapters 1 to 3
    print("Writing Chapters 1 to 3...")
    lines_13 = chapters_1_3.split("\n")
    paragraph_buffer = []
    
    for line in lines_13:
        stripped = line.strip()
        if not stripped:
            if paragraph_buffer:
                p_text = " ".join(paragraph_buffer)
                p = doc_full.add_paragraph(p_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)
                paragraph_buffer = []
            continue
            
        # Check if it's a heading
        is_heading = False
        if stripped.startswith("CHAPTER") or stripped in ["INTRODUCTION", "LITERATURE REVIEW", "METHODOLOGY"]:
            is_heading = True
            level = 1
        elif re.match(r'^[1-3]\.[0-9]+(\.[0-9]+)?\s+[A-Z]', stripped):
            is_heading = True
            level = 2 if len(stripped.split()[0].split('.')) == 2 else 3
            
        if is_heading:
            if paragraph_buffer:
                p_text = " ".join(paragraph_buffer)
                p = doc_full.add_paragraph(p_text)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)
                paragraph_buffer = []
                
            h = doc_full.add_heading(level=level)
            r = h.add_run(stripped)
            r.font.name = 'Arial'
            r.font.bold = True
            if level == 1:
                r.font.size = Pt(18)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                h.paragraph_format.space_before = Pt(24)
                h.paragraph_format.space_after = Pt(12)
                doc_full.add_page_break()
            elif level == 2:
                r.font.size = Pt(14)
                h.paragraph_format.space_before = Pt(18)
                h.paragraph_format.space_after = Pt(8)
            else:
                r.font.size = Pt(12)
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)
        else:
            # If the line ends with hyphen or normal letter, combine.
            paragraph_buffer.append(stripped)
            
    if paragraph_buffer:
        p_text = " ".join(paragraph_buffer)
        p = doc_full.add_paragraph(p_text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(6)
        
    # Add Chapter 4 & 5
    print("Writing Chapters 4 and 5...")
    parse_markdown_to_docx(doc_full, "Chapter_Four.md")
    parse_markdown_to_docx(doc_full, "Chapter_Five.md")
    
    # Add References
    print("Writing References...")
    doc_full.add_page_break()
    h = doc_full.add_heading(level=1)
    r = h.add_run("REFERENCES")
    r.font.name = 'Arial'
    r.font.size = Pt(18)
    r.font.bold = True
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.paragraph_format.space_before = Pt(24)
    h.paragraph_format.space_after = Pt(12)
    
    if references_text:
        ref_lines = references_text.split("\n")
        ref_buffer = []
        for line in ref_lines:
            stripped = line.strip()
            if not stripped:
                if ref_buffer:
                    p_text = " ".join(ref_buffer)
                    p = doc_full.add_paragraph(p_text)
                    p.paragraph_format.left_indent = Inches(0.5) # Hanging indent for APA references
                    p.paragraph_format.first_line_indent = Inches(-0.5)
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(6)
                    ref_buffer = []
                continue
            ref_buffer.append(stripped)
            
        if ref_buffer:
            p_text = " ".join(ref_buffer)
            p = doc_full.add_paragraph(p_text)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            
    doc_full.save("KIU_Student_Monitoring_Full_Thesis.docx")
    print("Saved KIU_Student_Monitoring_Full_Thesis.docx")

if __name__ == "__main__":
    compile_all_chapters()
